"""Exporters PDF / DOCX / HTML-Slides pour les livrables VITA-FORM."""
import io
import re
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
)
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import arabic_reshaper
from bidi.algorithm import get_display


GOLD = colors.HexColor("#B8860B")
DEEP_BLUE = colors.HexColor("#0A1128")
SLATE = colors.HexColor("#334155")

ARABIC_FONT_PATH = "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"
ARABIC_FONT_BOLD_PATH = "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf"
_ARABIC_FONTS_REGISTERED = False


def _ensure_arabic_fonts() -> None:
    global _ARABIC_FONTS_REGISTERED
    if _ARABIC_FONTS_REGISTERED:
        return
    try:
        pdfmetrics.registerFont(TTFont("NotoArabic", ARABIC_FONT_PATH))
        pdfmetrics.registerFont(TTFont("NotoArabic-Bold", ARABIC_FONT_BOLD_PATH))
        _ARABIC_FONTS_REGISTERED = True
    except Exception:
        # Fallback: keep Latin fonts; Arabic glyphs may not render but rest works.
        pass


def _shape_ar(text: str) -> str:
    """Reshape Arabic text + apply BiDi for correct RTL glyph order."""
    if not text:
        return text
    try:
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    except Exception:
        return text


def _md_to_blocks(md: str):
    """Très simple parseur Markdown -> liste de tuples (type, text)."""
    blocks = []
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            blocks.append(("space", ""))
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("> "):
            blocks.append(("quote", line[2:].strip()))
        elif re.match(r"^\s*[-*]\s+", line):
            blocks.append(("li", re.sub(r"^\s*[-*]\s+", "", line)))
        elif re.match(r"^\s*\d+\.\s+", line):
            blocks.append(("oli", re.sub(r"^\s*\d+\.\s+", "", line)))
        else:
            blocks.append(("p", line.strip()))
    return blocks


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
    return text


def render_pdf(title: str, author: str, institution: str, content_md: str,
               language: str = "fr") -> bytes:
    is_ar = language == "ar"
    if is_ar:
        _ensure_arabic_fonts()

    body_font = "NotoArabic" if is_ar else "Times-Roman"
    bold_font = "NotoArabic-Bold" if is_ar else "Times-Bold"
    italic_font = "NotoArabic" if is_ar else "Times-Italic"
    body_align = TA_RIGHT if is_ar else TA_JUSTIFY

    def t(text: str) -> str:
        return _shape_ar(text) if is_ar else text

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=title, author="VITA-FORM",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=bold_font,
                       fontSize=22, textColor=DEEP_BLUE, spaceAfter=18, leading=26,
                       alignment=body_align)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=bold_font,
                       fontSize=16, textColor=GOLD, spaceBefore=14, spaceAfter=8,
                       alignment=body_align)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=bold_font,
                       fontSize=13, textColor=SLATE, spaceBefore=10, spaceAfter=6,
                       alignment=body_align)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=body_font,
                         fontSize=11, leading=16, alignment=body_align, spaceAfter=6)
    quote = ParagraphStyle("Quote", parent=body, leftIndent=20, rightIndent=20,
                          textColor=GOLD, fontName=italic_font, borderPadding=6)
    li = ParagraphStyle("Li", parent=body, leftIndent=18, bulletIndent=8)
    cover_title = ParagraphStyle("Cover", parent=styles["Title"], fontName=bold_font,
                                 fontSize=28, textColor=DEEP_BLUE, alignment=TA_CENTER,
                                 spaceAfter=12, leading=34)
    cover_sub = ParagraphStyle("CoverSub", parent=styles["Italic"], fontName=italic_font,
                              fontSize=14, textColor=GOLD, alignment=TA_CENTER, spaceAfter=8)

    story = []
    # Cover
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph(t("VITA-FORM"), cover_sub))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(t(title), cover_title))
    story.append(Spacer(1, 0.5 * cm))
    if is_ar:
        story.append(Paragraph(t("منصة بيداغوجية حيوية"), cover_sub))
        story.append(Paragraph(t(f"المؤسسة المستهدفة : {institution}"), body))
        story.append(Paragraph(t(f"المؤلف : {author}"), body))
        story.append(Paragraph(t(f"تاريخ الإصدار : {datetime.now().strftime('%d/%m/%Y')}"), body))
    else:
        story.append(Paragraph("Plateforme Pédagogique Vitaliste", cover_sub))
        story.append(Paragraph(f"Institution destinataire : {institution}", body))
        story.append(Paragraph(f"Auteur du parcours : {author}", body))
        story.append(Paragraph(f"Date d'émission : {datetime.now().strftime('%d %B %Y')}", body))
    story.append(Spacer(1, 6 * cm))
    if is_ar:
        story.append(Paragraph(
            t("« لا نتلاعب بالأرقام، بل بالأرواح. »") +
            "<br/>" + t("— النظرية الحيوية للمالية العامة"),
            quote))
    else:
        story.append(Paragraph(
            "« On ne manipule pas des chiffres, mais des âmes. »<br/>— Théorie Vitaliste des Finances Publiques",
            quote))
    story.append(PageBreak())

    for kind, txt in _md_to_blocks(content_md):
        if kind == "space":
            story.append(Spacer(1, 0.2 * cm))
        elif kind == "h1":
            story.append(Paragraph(t(_strip_md_inline(txt)), h1))
        elif kind == "h2":
            story.append(Paragraph(t(_strip_md_inline(txt)), h2))
        elif kind == "h3":
            story.append(Paragraph(t(_strip_md_inline(txt)), h3))
        elif kind == "quote":
            story.append(Paragraph(t(_strip_md_inline(txt)), quote))
        elif kind in ("li", "oli"):
            bullet = "• "
            if is_ar:
                story.append(Paragraph(t(_strip_md_inline(txt)) + " " + bullet, li))
            else:
                story.append(Paragraph(f"{bullet}{_strip_md_inline(txt)}", li))
        else:
            story.append(Paragraph(t(_strip_md_inline(txt)), body))

    doc.build(story)
    return buffer.getvalue()


def _set_paragraph_rtl(paragraph) -> None:
    """Apply RTL directionality on a docx paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def render_docx(title: str, author: str, institution: str, content_md: str,
                language: str = "fr") -> bytes:
    is_ar = language == "ar"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = document.styles["Normal"]
    style.font.name = "Arial" if is_ar else "Calibri"
    style.font.size = Pt(11)
    if is_ar:
        # Arabic complex script font
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:cs"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")

    # Cover
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VITA-FORM")
    run.italic = True
    run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
    run.font.size = Pt(14)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(title_p)
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x0A, 0x11, 0x28)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(meta)
        meta.add_run(
            f"\nمنصة بيداغوجية حيوية\nالمؤسسة : {institution}\n"
            f"المؤلف : {author}\nالتاريخ : {datetime.now().strftime('%d/%m/%Y')}"
        )
    else:
        meta.add_run(
            f"\nPlateforme Pédagogique Vitaliste\nInstitution : {institution}\n"
            f"Auteur : {author}\nDate : {datetime.now().strftime('%d/%m/%Y')}"
        )

    document.add_paragraph()
    quote = document.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(quote)
        qrun = quote.add_run("« لا نتلاعب بالأرقام، بل بالأرواح. »")
    else:
        qrun = quote.add_run("« On ne manipule pas des chiffres, mais des âmes. »")
    qrun.italic = True
    qrun.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    document.add_page_break()

    for kind, txt in _md_to_blocks(content_md):
        clean = re.sub(r"\*\*|\*|`", "", txt)
        if kind == "space":
            document.add_paragraph()
            continue
        if kind == "h1":
            par = document.add_heading(clean, level=1)
        elif kind == "h2":
            par = document.add_heading(clean, level=2)
        elif kind == "h3":
            par = document.add_heading(clean, level=3)
        elif kind == "quote":
            par = document.add_paragraph(style="Intense Quote")
            par.add_run(clean).italic = True
        elif kind in ("li", "oli"):
            par = document.add_paragraph(clean, style="List Bullet")
        else:
            par = document.add_paragraph(clean)
        if is_ar:
            _set_paragraph_rtl(par)
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def render_slides_html(title: str, author: str, institution: str, content_md: str,
                       language: str = "fr") -> bytes:
    """Génère un fichier HTML autonome avec slides imprimables (1 slide / page)."""
    is_ar = language == "ar"
    blocks = _md_to_blocks(content_md)
    slides = []
    current = []
    for kind, txt in blocks:
        if kind in ("h1", "h2") and current:
            slides.append(current)
            current = []
        current.append((kind, txt))
    if current:
        slides.append(current)

    def render_block(kind, txt):
        clean = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", txt)
        clean = re.sub(r"\*(.+?)\*", r"<em>\1</em>", clean)
        if kind == "h1":
            return f"<h1>{clean}</h1>"
        if kind == "h2":
            return f"<h2>{clean}</h2>"
        if kind == "h3":
            return f"<h3>{clean}</h3>"
        if kind == "quote":
            return f"<blockquote>{clean}</blockquote>"
        if kind in ("li", "oli"):
            return f"<li>{clean}</li>"
        if kind == "space":
            return ""
        return f"<p>{clean}</p>"

    slides_html = []
    if is_ar:
        cover_tag = "VITA-FORM"
        cover_meta_inst = institution
        cover_meta_author = f"بقلم {author}"
        cover_quote = "« لا نتلاعب بالأرقام، بل بالأرواح. »"
    else:
        cover_tag = "VITA-FORM"
        cover_meta_inst = institution
        cover_meta_author = f"par {author}"
        cover_quote = "« On ne manipule pas des chiffres, mais des âmes. »"

    cover = f"""
    <section class="slide cover">
      <div class="cover-tag">{cover_tag}</div>
      <h1>{title}</h1>
      <p class="cover-meta">{cover_meta_inst}</p>
      <p class="cover-meta">{cover_meta_author}</p>
      <p class="cover-quote">{cover_quote}</p>
    </section>
    """
    slides_html.append(cover)
    for slide in slides:
        body = "\n".join(render_block(k, t) for k, t in slide if t.strip())
        slides_html.append(f'<section class="slide">{body}</section>')

    if is_ar:
        font_import = (
            "<link rel='preconnect' href='https://fonts.googleapis.com'>"
            "<link href='https://fonts.googleapis.com/css2?family=Amiri:wght@400;700&"
            "family=Noto+Naskh+Arabic:wght@400;700&display=swap' rel='stylesheet'>"
        )
        body_font = "'Noto Naskh Arabic', 'Amiri', serif"
        lang_attr = "ar"
        dir_attr = "rtl"
    else:
        font_import = ""
        body_font = "Georgia, 'Times New Roman', serif"
        lang_attr = "fr"
        dir_attr = "ltr"

    html = f"""<!doctype html>
<html lang=\"{lang_attr}\" dir=\"{dir_attr}\"><head><meta charset="utf-8"><title>{title} — VITA-FORM</title>
{font_import}
<style>
@page {{ size: A4 landscape; margin: 0; }}
body {{ margin:0; font-family: {body_font}; background:#0A1128; color:#F8FAFC; }}
.slide {{ page-break-after: always; width:100%; min-height: 100vh; padding: 4rem 5rem;
         background: linear-gradient(135deg, #0A1128 0%, #131B33 100%);
         box-sizing: border-box; }}
[dir="ltr"] .slide {{ border-left: 12px solid #D4AF37; }}
[dir="rtl"] .slide {{ border-right: 12px solid #D4AF37; }}
.slide.cover {{ display:flex; flex-direction:column; justify-content:center; align-items:center; text-align:center; }}
.cover-tag {{ letter-spacing: .3em; color:#D4AF37; font-size: 1rem; text-transform:uppercase; margin-bottom: 2rem; }}
.cover h1 {{ font-size: 3.5rem; margin: 0 0 1rem; color:#F8FAFC; }}
.cover-meta {{ color:#94A3B8; font-style: italic; margin: .25rem 0; }}
.cover-quote {{ margin-top: 4rem; color:#D4AF37; font-style: italic; }}
h1 {{ color:#F8FAFC; font-size: 2.4rem; border-bottom: 2px solid #D4AF37; padding-bottom: .5rem; }}
h2 {{ color:#D4AF37; font-size: 1.9rem; }}
h3 {{ color:#F3E5AB; font-size: 1.4rem; }}
p, li {{ font-size: 1.1rem; line-height: 1.8; color:#E2E8F0; }}
[dir="ltr"] blockquote {{ border-left: 4px solid #D4AF37; padding-left: 1rem; color:#F3E5AB; font-style: italic; }}
[dir="rtl"] blockquote {{ border-right: 4px solid #D4AF37; padding-right: 1rem; color:#F3E5AB; font-style: italic; }}
@media print {{ .slide {{ min-height: auto; height: 100vh; }} }}
</style></head><body>
{''.join(slides_html)}
</body></html>"""
    return html.encode("utf-8")
